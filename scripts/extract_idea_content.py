"""Liest pro Idee Anhang + Beschreibung aus edu-sharing, extrahiert Klartext
und sammelt das Ergebnis als JSON für anschließende Metadaten-Anreicherung.

Aufruf:
    python scripts/extract_idea_content.py [--limit 5] [--out content.json]
"""
from __future__ import annotations
import argparse
import asyncio
import json
import os
import sys
from base64 import b64encode
from io import BytesIO

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

import httpx
from app.db import connect


def _read_pdf_text(blob: bytes, max_chars: int = 6000) -> str:
    try:
        import pdfplumber
        out: list[str] = []
        with pdfplumber.open(BytesIO(blob)) as pdf:
            for page in pdf.pages[:8]:  # max 8 Seiten
                t = page.extract_text() or ""
                out.append(t)
                if sum(len(x) for x in out) > max_chars:
                    break
        return "\n".join(out)[:max_chars]
    except Exception as e:
        return f"[PDF-Extract-Error: {e}]"


def _read_html_text(blob: bytes, max_chars: int = 4000) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(blob, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.decompose()
        text = soup.get_text("\n", strip=True)
        return text[:max_chars]
    except Exception as e:
        return f"[HTML-Extract-Error: {e}]"


def _read_docx_text(blob: bytes, max_chars: int = 6000) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(blob))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t: parts.append(t)
        # Tabellen-Inhalte mit aufnehmen — Steckbriefe nutzen viel Tabellen
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return text[:max_chars] if text else "[DOCX-leer-keine-Paragraphen]"
    except Exception as e:
        return f"[DOCX-Extract-Error: {e}]"


def _extract_text(name: str, mime: str, blob: bytes) -> str:
    name_lc = (name or "").lower()
    mime_lc = (mime or "").lower()
    if "pdf" in mime_lc or name_lc.endswith(".pdf"):
        return _read_pdf_text(blob)
    if "html" in mime_lc or name_lc.endswith(".html") or name_lc.endswith(".htm"):
        return _read_html_text(blob)
    if name_lc.endswith(".docx") or "officedocument.wordprocessingml" in mime_lc:
        return _read_docx_text(blob)
    if mime_lc.startswith("text/") or name_lc.endswith(".txt") or name_lc.endswith(".md"):
        try:
            return blob.decode("utf-8", errors="replace")[:6000]
        except Exception:
            return "[plain-decode-error]"
    return f"[no-extractor-for: name={name} mime={mime} size={len(blob)}]"


async def _fetch_node(client: httpx.AsyncClient, repo: str, node_id: str, auth: str) -> dict:
    r = await client.get(
        f"{repo}/node/v1/nodes/-home-/{node_id}/metadata?propertyFilter=-all-",
        headers={"Authorization": auth}, timeout=30,
    )
    r.raise_for_status()
    return r.json().get("node") or {}


async def _fetch_content(client: httpx.AsyncClient, url: str, auth: str) -> bytes:
    r = await client.get(url, headers={"Authorization": auth}, timeout=60)
    r.raise_for_status()
    return r.content


async def main(limit: int | None, outfile: str) -> None:
    # Credentials aus .env
    env = os.path.join(os.path.dirname(__file__), "..", ".env")
    es_user = es_pass = None
    if os.path.exists(env):
        for line in open(env, encoding="utf-8"):
            line = line.strip()
            if line.startswith("EDU_GUEST_USER="):
                es_user = line.split("=", 1)[1].strip()
            elif line.startswith("EDU_GUEST_PASS="):
                es_pass = line.split("=", 1)[1].strip()
    if not (es_user and es_pass):
        print("Keine ES-Credentials in .env gefunden")
        sys.exit(2)
    auth = "Basic " + b64encode(f"{es_user}:{es_pass}".encode()).decode()
    repo = "https://redaktion.openeduhub.net/edu-sharing/rest"

    with connect() as con:
        rows = con.execute(
            """SELECT id,title,description,attachment_name,attachment_mimetype,
                      attachment_size,attachment_url,phase,events,keywords,topic_id
                 FROM idea ORDER BY created_at DESC"""
        ).fetchall()
    if limit: rows = rows[:limit]

    out: list[dict] = []
    async with httpx.AsyncClient() as client:
        for i, r in enumerate(rows, 1):
            entry = {
                "id": r["id"], "title": r["title"],
                "old_description": r["description"] or "",
                "old_keywords": r["keywords"] or "",
                "old_phase": r["phase"], "old_events": r["events"] or "",
                "topic_id": r["topic_id"],
                "file_name": r["attachment_name"],
                "file_mime": r["attachment_mimetype"],
                "file_size": r["attachment_size"] or 0,
                "extracted_text": "",
                "preview_url": None,
                "extract_error": None,
            }
            print(f"[{i}/{len(rows)}] {r['id'][:8]}  {(r['attachment_name'] or '-')[:40]} ({r['attachment_size']} B)", flush=True)
            try:
                node = await _fetch_node(client, repo, r["id"], auth)
                preview = (node.get("preview") or {}).get("url")
                entry["preview_url"] = preview
                # 1) Wenn Knoten ein Reference ist: Original benutzen
                target_id = node.get("originalId") or r["id"]
                # 2) /textContent liefert für HTML-Brainstorm-Karten das raw-HTML
                #    direkt als JSON. Das ist die zuverlässigste Quelle für die
                #    kleinen Stubs, die kein Download-Asset haben.
                #    Achtung: für PDFs/DOCX liefert der Endpoint BLOB-Bytes als
                #    raw-Feld — daher nur für HTML/Text-MIMEs nutzen.
                text_from_textcontent: str | None = None
                mime = (r["attachment_mimetype"] or "").lower()
                is_html_like = ("html" in mime or "text/" in mime
                                or (r["attachment_name"] or "").lower().endswith((".html",".htm",".txt",".md")))
                if is_html_like:
                    try:
                        rt = await client.get(
                            f"{repo}/node/v1/nodes/-home-/{target_id}/textContent",
                            headers={"Authorization": auth}, timeout=30,
                        )
                        if rt.status_code == 200:
                            j = rt.json() or {}
                            raw = j.get("raw") or j.get("html") or j.get("text") or ""
                            if raw:
                                text_from_textcontent = _read_html_text(raw.encode("utf-8"))
                    except Exception:
                        pass
                # 3) Fallback: Bytes via downloadUrl/content-Endpoint
                blob: bytes | None = None
                if not text_from_textcontent:
                    for url in (node.get("downloadUrl"), r["attachment_url"],
                                f"{repo}/node/v1/nodes/-home-/{target_id}/content"):
                        if not url: continue
                        try:
                            blob = await _fetch_content(client, url, auth)
                            if blob: break
                        except Exception:
                            continue
                if text_from_textcontent:
                    entry["extracted_text"] = text_from_textcontent
                elif blob:
                    entry["extracted_text"] = _extract_text(
                        r["attachment_name"] or "", r["attachment_mimetype"] or "", blob,
                    )
                else:
                    entry["extract_error"] = "no-content-available"
            except httpx.HTTPStatusError as e:
                entry["extract_error"] = f"HTTP {e.response.status_code}"
            except Exception as e:
                entry["extract_error"] = str(e)[:200]
            out.append(entry)

    with open(outfile, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print(f"\nGeschrieben: {outfile}  ({len(out)} Einträge)")


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--limit", type=int, default=None)
    p.add_argument("--out", default="ideen_content.json")
    args = p.parse_args()
    asyncio.run(main(args.limit, args.out))
